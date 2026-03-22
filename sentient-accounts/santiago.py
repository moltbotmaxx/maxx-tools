from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import NoSuchElementException, StaleElementReferenceException
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
import re
from datetime import datetime

# ─────────────────────────────────────────────
# CONFIGURACIÓN
# ─────────────────────────────────────────────
CHROME_PROFILE = r"C:\Users\Santi\chrome_selenium_profile"
MAX_REELS      = 100
OUTPUT_EXCEL   = "instagram_reels.xlsx"
OUTPUT_DOCX    = "instagram_reels_report.docx"

# ─────────────────────────────────────────────
# DRIVER
# ─────────────────────────────────────────────
options = Options()
options.add_argument(f"--user-data-dir={CHROME_PROFILE}")
options.add_argument("--disable-notifications")
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")

driver = webdriver.Chrome(options=options)
driver.get("https://www.instagram.com")

input("\n[→] Abre el perfil y la pestaña de REELS. Luego presiona ENTER...\n")
time.sleep(3)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def parse_number(text):
    if not text:
        return None
    text = text.strip().lower()
    text = re.sub(r"(\d)\.(\d{3})", r"\1\2", text)
    text = text.replace(",", ".")
    match = re.search(r"[\d.]+", text)
    if not match:
        return None
    number = float(match.group())
    if "mil" in text or "k" in text:
        return int(number * 1_000)
    if "m" in text:
        return int(number * 1_000_000)
    return int(number)

def format_number(n):
    if n is None:
        return "N/D"
    if n >= 1_000_000:
        return f"{n/1_000_000:.2f}M"
    if n >= 1_000:
        return f"{n/1_000:.1f}K"
    return str(n)

# ─────────────────────────────────────────────
# GENERAR ARCHIVOS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_styled_table(doc, headers, rows, col_widths, header_color="4267B2"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = table.rows[0]
    for j, (cell, header) in enumerate(zip(hdr_row.cells, headers)):
        set_cell_bg(cell, header_color)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9)
        cell.width = Cm(col_widths[j])
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        bg  = "F0F4FF" if i % 2 == 0 else "FFFFFF"
        for j, (cell, value) in enumerate(zip(row.cells, row_data)):
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value) if value is not None else "N/D")
            run.font.size = Pt(9)
            cell.width = Cm(col_widths[j])
    return table

def guardar_archivos(data):
    if not data:
        print("\n⚠️  No hay datos para guardar (0 reels procesados).")
        return

    df       = pd.DataFrame(data)
    df_excel = df[["N°", "link", "views", "likes", "likes/views (%)"]].copy()
    df_excel.to_excel(OUTPUT_EXCEL, index=False)
    print(f"\n✅ Excel guardado: {OUTPUT_EXCEL}")

    numeric = df.dropna(subset=["views", "likes"])
    total   = len(df)
    valid   = len(numeric)

    stats      = {}
    top5       = None
    top5_ratio = None

    if valid > 0:
        stats = {
            "total_reels":    total,
            "reels_con_data": valid,
            "total_views":    int(numeric["views"].sum()),
            "avg_views":      int(numeric["views"].mean()),
            "max_views":      int(numeric["views"].max()),
            "min_views":      int(numeric["views"].min()),
            "total_likes":    int(numeric["likes"].sum()),
            "avg_likes":      int(numeric["likes"].mean()),
            "max_likes":      int(numeric["likes"].max()),
            "avg_ratio":      round(numeric["likes/views (%)"].mean(), 2) if numeric["likes/views (%)"].notna().any() else 0,
            "best_ratio":     round(numeric["likes/views (%)"].max(), 2) if numeric["likes/views (%)"].notna().any() else 0,
        }
        top5 = numeric.nlargest(min(5, valid), "views")[
            ["N°", "link", "views_fmt", "likes_fmt", "likes/views (%)"]
        ].reset_index(drop=True)
        top5_ratio = numeric.dropna(subset=["likes/views (%)"]).nlargest(min(5, valid), "likes/views (%)")[
            ["N°", "link", "views_fmt", "likes_fmt", "likes/views (%)"]
        ].reset_index(drop=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("📊 Reporte de Reels de Instagram")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x42, 0x67, 0xB2)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}  ·  Reels procesados: {total}")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    if stats:
        h = doc.add_paragraph()
        rh = h.add_run("Resumen General")
        rh.bold = True; rh.font.size = Pt(13); rh.font.color.rgb = RGBColor(0x42, 0x67, 0xB2)
        doc.add_paragraph()
        add_styled_table(doc, ["Métrica", "Valor"], [
            ["Reels analizados",            f"{stats['reels_con_data']} / {stats['total_reels']}"],
            ["Total de vistas",             format_number(stats["total_views"])],
            ["Promedio de vistas por reel", format_number(stats["avg_views"])],
            ["Reel con más vistas",         format_number(stats["max_views"])],
            ["Reel con menos vistas",       format_number(stats["min_views"])],
            ["Total de likes",              format_number(stats["total_likes"])],
            ["Promedio de likes por reel",  format_number(stats["avg_likes"])],
            ["Ratio likes/vistas promedio", f"{stats['avg_ratio']}%"],
            ["Mejor ratio likes/vistas",    f"{stats['best_ratio']}%"],
        ], col_widths=[8, 5])
        doc.add_paragraph()

        h2 = doc.add_paragraph()
        rh2 = h2.add_run("🔥 Top Reels por Vistas")
        rh2.bold = True; rh2.font.size = Pt(13); rh2.font.color.rgb = RGBColor(0x42, 0x67, 0xB2)
        doc.add_paragraph()
        add_styled_table(doc, ["N°", "Vistas", "Likes", "Ratio (%)"],
            [[r["N°"], r["views_fmt"], r["likes_fmt"],
              f"{r['likes/views (%)']:.2f}%" if r["likes/views (%)"] else "N/D"]
             for _, r in top5.iterrows()],
            col_widths=[1.5, 3, 3, 3], header_color="E44D26")
        doc.add_paragraph()

        h3 = doc.add_paragraph()
        rh3 = h3.add_run("💡 Top Reels por Engagement (likes/vistas)")
        rh3.bold = True; rh3.font.size = Pt(13); rh3.font.color.rgb = RGBColor(0x42, 0x67, 0xB2)
        doc.add_paragraph()
        add_styled_table(doc, ["N°", "Vistas", "Likes", "Ratio (%)"],
            [[r["N°"], r["views_fmt"], r["likes_fmt"], f"{r['likes/views (%)']:.2f}%"]
             for _, r in top5_ratio.iterrows()],
            col_widths=[1.5, 3, 3, 3], header_color="27AE60")
        doc.add_paragraph()

    h4 = doc.add_paragraph()
    rh4 = h4.add_run("📋 Todos los Reels")
    rh4.bold = True; rh4.font.size = Pt(13); rh4.font.color.rgb = RGBColor(0x42, 0x67, 0xB2)
    doc.add_paragraph()
    add_styled_table(doc, ["N°", "Vistas", "Likes", "Ratio (%)"],
        [[r["N°"], r["views_fmt"], r["likes_fmt"],
          f"{r['likes/views (%)']:.2f}%" if r["likes/views (%)"] is not None else "N/D"]
         for _, r in df.iterrows()],
        col_widths=[1.5, 3.5, 3.5, 3])

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run(f"⚠️  Los links completos están en el archivo Excel: {OUTPUT_EXCEL}")
    rn.font.size = Pt(8); rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT_DOCX)
    print(f"✅ Reporte Word guardado: {OUTPUT_DOCX}")
    print(f"   → Súbelo a Google Docs: drive.google.com → Nuevo → Subir archivo")

# ─────────────────────────────────────────────
# EJECUCIÓN PRINCIPAL — scroll + scraping juntos
# ─────────────────────────────────────────────
data        = []
errors      = 0
procesados  = set()   # links ya procesados para no repetir
actions     = ActionChains(driver)

print("Iniciando scroll y scraping simultáneo...")
print("(Presiona Ctrl+C en cualquier momento para parar y guardar)\n")

try:
    scroll_sin_nuevos = 0

    while len(data) < MAX_REELS and scroll_sin_nuevos < 6:

        # Buscar todos los reels visibles en pantalla
        reels_visibles = driver.find_elements(By.XPATH, "//a[contains(@href,'/reel/')]")
        nuevos = [r for r in reels_visibles if r.get_attribute("href") not in procesados]

        if not nuevos:
            scroll_sin_nuevos += 1
            driver.execute_script("window.scrollBy(0, 600);")
            time.sleep(0.8)
            continue

        scroll_sin_nuevos = 0

        # Procesar cada reel nuevo encontrado
        for reel in nuevos:
            if len(data) >= MAX_REELS:
                break

            link = views_text = likes_text = None

            try:
                link = reel.get_attribute("href")
                if not link or link in procesados:
                    continue
                procesados.add(link)

                # Vistas
                try:
                    full_text = reel.text.strip()
                    if full_text:
                        views_text = full_text.split("\n")[-1]
                except StaleElementReferenceException:
                    pass

                # Hover para likes
                actions.move_to_element(reel).perform()
                time.sleep(0.5)

                try:
                    span = reel.find_element(By.XPATH, ".//span[contains(@class,'x')]")
                    likes_text = span.text
                except NoSuchElementException:
                    try:
                        span = reel.find_element(By.XPATH, ".//span")
                        likes_text = span.text
                    except NoSuchElementException:
                        pass

                views = parse_number(views_text)
                likes = parse_number(likes_text)

                if views and likes and likes > views:
                    views, likes = likes, views

                ratio = round((likes / views) * 100, 2) if views and likes else None

                data.append({
                    "N°": len(data) + 1,
                    "link": link,
                    "views": views,
                    "likes": likes,
                    "likes/views (%)": ratio,
                    "views_fmt": format_number(views),
                    "likes_fmt": format_number(likes),
                })

                print(f"  [{len(data):>3}/{MAX_REELS}] "
                      f"{format_number(views):>8} vistas  |  "
                      f"{format_number(likes):>8} likes")

            except StaleElementReferenceException:
                continue
            except Exception:
                errors += 1
                if link:
                    procesados.add(link)
                continue

        # Scroll para cargar más
        driver.execute_script("window.scrollBy(0, 600);")
        time.sleep(0.8)

except KeyboardInterrupt:
    print(f"\n\n⏹️  Pausado. Reels procesados: {len(data)}")
    print("Guardando archivos...\n")

finally:
    try:
        driver.quit()
    except Exception:
        pass
    guardar_archivos(data)
    if errors:
        print(f"\n⚠️  {errors} reels tuvieron errores al procesar.")